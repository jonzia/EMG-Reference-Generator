# Import libraries
import tkinter as tk
from tkinter.filedialog import askopenfilename
from os.path import dirname
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt

# GLOBAL VARIABLES
isUILocked = False
filename = "filename.docx"
age, height = 0, 0

# Dictionaries of normal values
nerveConduction = {
    "medianSNAP": [14, 13, 11, 9, 6],
    "medianSNCV": [48, 47, 45, 42, 40],
    "medianCMAP": [6, 6, 5, 4.5, 4.5],
    "medianDML": [4, 4.2, 4.3, 4.5, 4.5],
    "medianMNCV": [48, 47, 45, 42, 40],
    "ulnarSNAP": [12, 11, 9, 7, 5],
    "ulnarSNCV": [47, 45, 42, 40, 38],
    "ulnarCMAP": [5.5, 5.5, 4.8, 4.5, 4.5],
    "ulnarDML": [3.7, 3.9, 4.0, 4.2, 4.2],
    "ulnarMNCVarm": [48, 47, 45, 42, 40],
    "ulnarMNCVelb": [46, 45, 42, 40, 40],
    "radialSNAP": [12, 11, 10, 9, 8],
    "radialSNCV": [50, 48, 47, 44, 42],
    "musculoSNAP": [12, 10, 9, 7, 5],
    "musculoSNCV": [50, 48, 47, 44, 42],
    "suralSNAP": [9, 7, 5, 2, 0],
    "suralSNCV": [40, 38, 36, 35, 32],
    "fibularSNAP": [4, 4, 2, 0, 0],
    "fibularSNCV": [40, 38, 36, 35, 32],
    "fibularDML": [5.2, 5.5, 5.5, 5.8, 5.8],
    "fibularCMAP": [2.5, 2, 2, 1.5, 1.5],
    "fibularMNCVleg": [42, 40, 38, 36, 34],
    "fibularMNCVhead": [40, 38, 36, 34, 32],
    "tibialCMAP": [4.5, 4.5, 3, 2, 2],
    "tibialMNCV": [42, 40, 38, 36, 34]
}

fWaveLatency = {
    "median": [[24.1, 24.9, 25.6, 26.4, 27.1, 27.9, 28.7, 29.4, 30.2], [25.1, 25.9, 26.6, 27.4, 28.1, 28.9, 29.7, 30.4, 31.2], [26.1, 26.9, 27.6, 28.4, 29.1, 29.9, 30.7, 31.4, 32.2]],
    "ulnar": [[24.1, 25.0, 25.8, 26.7, 27.6, 28.4, 29.3, 30.2, 31.0], [25.1, 25.9, 26.8, 27.7, 28.5, 29.4, 30.3, 31.1, 32.0], [26.0, 26.9, 27.8, 28.6, 29.5, 30.4, 31.2, 32.1, 33.0]],
    "peroneal": [[47.8, 49.1, 50.4, 51.6, 52.9, 54.2, 55.5, 56.8, 58.1], [47.8, 49.1, 50.4, 51.6, 52.9, 54.2, 55.5, 56.8, 58.1], [47.8, 49.1, 50.4, 51.6, 52.9, 54.2, 55.5, 56.8, 58.1]],
    "tibial": [[50.0, 51.2, 52.4, 53.6, 54.8, 56.0, 57.2, 58.4, 59.6], [50.0, 51.2, 52.4, 53.6, 54.8, 56.0, 57.2, 58.4, 59.6], [50.0, 51.2, 52.4, 53.6, 54.8, 56.0, 57.2, 58.4, 59.6]]
}

nerveKeywords = "median ulnar radial musculo peroneal tibial"
levelKeywords = "wrist elbow axilla ankle head fossa"

# Function declaration
def getMotorValue(nerve, level, value, ncsIdx):
    keyword = nerve
        
    if value == 1:
        keyword += "DML"
    elif value == 2:
        keyword += "CMAP"
    elif value == 6:
        keyword += "MNCV"
        
    if keyword == "ulnarMNCV":
        if "elb" in level.lower():
            keyword = "ulnarMNCVelb"
        else:
            keyword = "ulnarMNCVarm"
            
    if keyword == "fibularMNCV":
        if "head" in level.lower():
            keyword = "fibularMNCVhead"
        else:
            keyword = "fibularMNCVleg"
            
    return nerveConduction[keyword][ncsIdx]

def getFWaveValue(nerve, fwIdx, fwHeightIdx):
    return fWaveLatency[nerve][fwIdx][fwHeightIdx]

def getSensoryValue(nerve, value, ncsIdx):
    keyword = nerve
    
    if value == 3:
        keyword += "SNAP"
    elif value == 5:
        keyword += "SNCV"
        
    return nerveConduction[keyword][ncsIdx]

def setFilename():
    global isUILocked, filename
    if isUILocked: return
    filename = askopenfilename()
    sourceText["text"] = "File Path: " + filename

def generateFile():
    global isUILocked, age, height, filename, success
    if isUILocked: return
    # If a file has not been selected, throw an error
    if filename == "filename.docx":
        progressText["text"] = "Please select a file."
        return
    # If there is no age defined, throw an error
    try:
        age = float(ageEntry.get())
    except:
        progressText["text"] = "Error: Please enter valid age."
        return
    # If there is no height, throw an error
    try:
        height = float(heightEntry.get())
    except:
        progressText["text"] = "Error: Please enter valid height."
        return
    # Lock the user interface
    isUILocked = True
    progressText["text"] = "Working..."
    # To Do
    success = runProgram()
    if success:
        progressText["text"] = "File Saved: " + dirname(filename) + "/emgref.docx"
    else:
        progressText["text"] = "Error generating file."
    isUILocked = False

def exitGUI():
    window.destroy()

def runProgram():
    global age, height, filename
    # try:
    # Determine proper list index for nerve conduction studies based on patient age
    # (0) 4-39, (1) 40-59, (2) 60-69, (3) 70-79, (4) 80+
    ncsIdx = 0 if age < 40 else 1 if age < 60 else 2 if age < 70 else 3 if age < 80 else 4

    # Determine proper list index for F-wave latency studies based on patient age and height
    # (0) 30-49, (1) 50-69, (2) 70+
    fwIdx = 0 if age < 50 else 1 if age < 70 else 2
    fwHeightIdx = 0 if height < 60 else 1 if height < 62 else 2 if height < 64 else 3 if height < 66 else 4 if height < 68 else 5 if height < 70 else 6 if height < 72 else 7 if height < 74 else 8

    # Import word document
    document = Document(filename)

    # Create a new document
    newDoc = Document()
    style = newDoc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(8)

    # document.tables will contain 3 tables:
    # 1. Motor nerve conduction studies
    # 2. F-wave
    # 3. Sensory nerve conduction studies

    # 1. Motor nerve conduction studies
    newDoc.add_paragraph("Motor Nerve Conduction Studies")

    # Populate the table
    t = document.tables[0]
    numRows = len(t.rows)
    numCols = len(t.columns)
    newTable = newDoc.add_table(numRows, 9)
    newTable.style = "Table Grid"
    # Populate header and first column
    for i in range(0,2):
        for j, column in enumerate(t.columns):
            newTable.cell(i,j).text = t.cell(i,j).text
    for i in range(2,numRows):
        newTable.cell(i,0).text = t.cell(i,0).text
    # Populate the rest of the table by row
    for i in range(2,numRows):
        # If a nerve keyword is present, merge the cells and skip to the next line
        FLAG = False
        for j, word in enumerate(newTable.cell(i,0).text.split(" ")):
            if word.lower() in nerveKeywords:
                newTable.cell(i,0).merge(newTable.cell(i,numCols-1))
                nerveKeyword = newTable.cell(i,0).text.split(" ")[j].lower()
                FLAG, FLAG2 = True, True
                break
        if FLAG:
            continue
        # If a level keyword is present, update the level keyword and loop through the columns
        FLAG = False
        for j, word in enumerate(newTable.cell(i,0).text.split(" ")):
            if word.lower() in levelKeywords:
                levelKeyword = newTable.cell(i,0).text.split(" ")[j].lower()
                for k in range(1,numCols):
                    if k == 1 or k == 2 or k == 6:
                        p = newTable.cell(i,k).add_paragraph()
                        value = getMotorValue(nerveKeyword, levelKeyword, k, ncsIdx)
                        # If the last row was a keyword, don't provide reference for 6
                        if k == 6 and FLAG2:
                            cellText = t.cell(i,k).text
                            runner = p.add_run(cellText)
                            FLAG2 = False
                            continue
                        # If k = 1, only provide reference if last row was a keyword
                        elif k == 1 and not FLAG2:
                            cellText = t.cell(i,k).text
                            runner = p.add_run(cellText)
                            continue
                        else:
                            cellText = t.cell(i,k).text + " (" + str(value) + ")"
                            runner = p.add_run(cellText)
                            font = runner.font
                            # Bold if abnormal
                            FLAG3 = False
                            try:
                                if k == 1 and float(t.cell(i,k).text) > value:
                                    FLAG3 = True
                                elif k == 2 and float(t.cell(i,k).text) < value:
                                    FLAG3 = True
                                elif k == 6 and float(t.cell(i,k).text) < value:
                                    FLAG3 = True
                                if FLAG3:
                                    font.color.rgb = RGBColor(255, 0, 0)
                                    runner.bold = True
                            except:
                                continue
                    else:
                        newTable.cell(i,k).text = t.cell(i,k).text

    # 2. F-wave
    newDoc.add_paragraph("F-Wave")

    # Populate the table
    t = document.tables[1]
    numRows = len(t.rows)
    numCols = len(t.columns)
    newTable = newDoc.add_table(numRows, 4)
    newTable.style = "Table Grid"
    # Populate header and first column
    for i in range(0,2):
        for j, column in enumerate(t.columns):
            newTable.cell(i,j).text = t.cell(i,j).text
    for i in range(2,numRows):
        newTable.cell(i,0).text = t.cell(i,0).text
    # Populate the rest of the table by row
    for i in range(2,numRows):
        # If a nerve keyword is present, merge the cells and skip to the next line
        FLAG = False
        for j, word in enumerate(newTable.cell(i,0).text.split(" ")):
            if word.lower() in nerveKeywords:
                newTable.cell(i,0).merge(newTable.cell(i,numCols-1))
                nerveKeyword = newTable.cell(i,0).text.split(" ")[j].lower()
                FLAG = True
                break
        if FLAG:
            continue
        # If a level keyword is present, update the level keyword and loop through the columns
        FLAG = False
        for j, word in enumerate(newTable.cell(i,0).text.split(" ")):
            if word.lower() in levelKeywords:
                levelKeyword = newTable.cell(i,0).text.split(" ")[j].lower()
                for k in range(1,numCols):
                    if k == 1:
                        p = newTable.cell(i,k).add_paragraph()
                        value = getFWaveValue(nerveKeyword, fwIdx, fwHeightIdx)
                        cellText = t.cell(i,k).text + " (" + str(value) + ")"
                        runner = p.add_run(cellText)
                        font = runner.font
                        # Bold if abnormal
                        try:
                            if float(t.cell(i,k).text) > value:
                                font.color.rgb = RGBColor(255, 0, 0)
                                runner.bold = True
                        except:
                            continue
                    else:
                        newTable.cell(i,k).text = t.cell(i,k).text

    # 3. Sensory nerve conduction studies
    newDoc.add_paragraph("Sensory Nerve Conduction Studies")

    # Populate the table
    t = document.tables[2]
    numRows = len(t.rows)
    numCols = len(t.columns)
    newTable = newDoc.add_table(numRows, 8)
    newTable.style = "Table Grid"
    # Populate header and first column
    for i in range(0,2):
        for j, column in enumerate(t.columns):
            newTable.cell(i,j).text = t.cell(i,j).text
    for i in range(2,numRows):
        newTable.cell(i,0).text = t.cell(i,0).text
    # Populate the rest of the table by row
    for i in range(2,numRows):
        # If a nerve keyword is present, merge the cells and skip to the next line
        FLAG = False
        for j, word in enumerate(newTable.cell(i,0).text.split(" ")):
            if word.lower() in nerveKeywords:
                newTable.cell(i,0).merge(newTable.cell(i,numCols-1))
                nerveKeyword = newTable.cell(i,0).text.split(" ")[j].lower()
                FLAG = True
                break
        if FLAG:
            continue
        # If a nerve keyword is not present, loop through the columns
        for k in range(1,numCols):
            if k == 3 or k == 5:
                p = newTable.cell(i,k).add_paragraph()
                value = getSensoryValue(nerveKeyword, k, ncsIdx)
                cellText = t.cell(i,k).text + " (" + str(value) + ")"
                runner = p.add_run(cellText)
                font = runner.font
                # Bold if abnormal
                try:
                    if float(t.cell(i,k).text) < value:
                        font.color.rgb = RGBColor(255, 0, 0)
                        runner.bold = True
                except:
                    continue
            else:
                newTable.cell(i,k).text = t.cell(i,k).text
                
    newDoc.save(dirname(filename) + "/emgref.docx")
    return True
    # except:
    #     return False

# Create GUI
window = tk.Tk()
window.title("EMG Reference Generator")

# GUI header
frame1 = tk.Frame(master = window)
frame1.pack()
titleText = tk.Label(master = frame1, text = "EMG Reference Generator")
titleText.pack()
authorText = tk.Label(master = frame1, text = "(c) 2024 Jon Zia")
authorText.pack()

# GUI buttons
frame2 = tk.Frame(master = window)
frame2.pack()
sourceButton = tk.Button(master = frame2, text = "Select File", command = setFilename)
sourceButton.grid(row = 0, column = 0)
exitButton = tk.Button(master = frame2, text = "Exit", command = exitGUI)
exitButton.grid(row = 0, column = 1)

# GUI data
frame3 = tk.Frame(master = window)
frame3.pack()
agePrompt = tk.Label(master = frame3, text = "Patient Age:")
agePrompt.grid(row = 0, column = 0)
heightPrompt = tk.Label(master = frame3, text = "Patient Height (in):")
heightPrompt.grid(row = 1, column = 0)
ageEntry = tk.Entry(master = frame3)
ageEntry.grid(row = 0, column = 1)
heightEntry = tk.Entry(master = frame3)
heightEntry.grid(row = 1, column = 1)


# GUI text output
frame4 = tk.Frame(master = window)
frame4.pack()
generateButton = tk.Button(master = frame4, text = "Generate", command = generateFile)
generateButton.pack()
sourceText = tk.Label(master = frame4, text = "File Path:")
sourceText.pack()
progressText = tk.Label(master = frame4, text = "")
progressText.pack()

window.mainloop()